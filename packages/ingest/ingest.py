"""File ingestion: turn a single upload into a list of pages.

Supports:
- PDF (rasterized via pypdfium2; text layer captured if present)
- Images (jpeg/png/tiff/bmp/webp — treated as a single-page document)
- DOCX (text-only; a single synthetic page with the extracted text)

Output is a pure-Python dataclass structure. Persistence and pipeline
orchestration happen elsewhere.
"""
from __future__ import annotations

import enum
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from PIL import Image


class SourceKind(str, enum.Enum):
    IMAGE = "image"
    PDF_SCANNED = "pdf_scanned"
    PDF_TEXT_LAYER = "pdf_text_layer"
    DOCX = "docx"
    UNKNOWN = "unknown"


@dataclass
class IngestedPage:
    page_index: int
    image_path: Optional[Path] = None
    text_layer: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None

    @property
    def has_image(self) -> bool:
        return self.image_path is not None

    @property
    def text_layer_used(self) -> bool:
        return bool(self.text_layer) and len(self.text_layer.strip()) >= 40


@dataclass
class IngestedDocument:
    source_path: Path
    kind: SourceKind
    pages: list[IngestedPage] = field(default_factory=list)
    note: Optional[str] = None


IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}

# pypdfium2 renders at 72 DPI by default; 2.0 = ~144 DPI which is a good
# balance between OCR quality and memory use.
PDF_RENDER_SCALE = 2.0


def ingest_file(
    source_path: Path,
    pages_dir: Path,
    *,
    mime_hint: Optional[str] = None,
) -> IngestedDocument:
    """Ingest a single file into pages.

    Parameters
    ----------
    source_path:
        Path to the uploaded file on disk.
    pages_dir:
        Directory where per-page PNG images should be written. Created if
        it does not exist.
    mime_hint:
        Optional mime-type hint from the HTTP upload, used as a fallback
        when the extension is ambiguous.
    """
    source_path = Path(source_path)
    pages_dir = Path(pages_dir)
    pages_dir.mkdir(parents=True, exist_ok=True)

    suffix = source_path.suffix.lower()
    if suffix in PDF_EXTS or (mime_hint == "application/pdf"):
        return _ingest_pdf(source_path, pages_dir)
    if suffix in IMAGE_EXTS or (mime_hint and mime_hint.startswith("image/")):
        return _ingest_image(source_path, pages_dir)
    if suffix in DOCX_EXTS or mime_hint in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _ingest_docx(source_path)

    return IngestedDocument(
        source_path=source_path,
        kind=SourceKind.UNKNOWN,
        pages=[],
        note=f"unsupported extension/mime {suffix!r}/{mime_hint!r}",
    )


# --- image ---------------------------------------------------------------


def _ingest_image(source_path: Path, pages_dir: Path) -> IngestedDocument:
    with Image.open(source_path) as img:
        img = img.convert("RGB")
        width, height = img.size
        out = pages_dir / "page_0001.png"
        img.save(out, format="PNG")
    return IngestedDocument(
        source_path=source_path,
        kind=SourceKind.IMAGE,
        pages=[IngestedPage(page_index=0, image_path=out, width=width, height=height)],
    )


# --- pdf -----------------------------------------------------------------


def _ingest_pdf(source_path: Path, pages_dir: Path) -> IngestedDocument:
    import pypdfium2 as pdfium  # lazy import — heavy

    pdf = pdfium.PdfDocument(str(source_path))
    pages: list[IngestedPage] = []
    try:
        any_text = False
        for i, page in enumerate(pdf):
            # render image
            pil = page.render(scale=PDF_RENDER_SCALE).to_pil().convert("RGB")
            width, height = pil.size
            out = pages_dir / f"page_{i + 1:04d}.png"
            pil.save(out, format="PNG")
            # try to grab text layer if present
            try:
                text_page = page.get_textpage()
                text = text_page.get_text_range()
                text_page.close()
            except Exception:
                text = ""
            if text and text.strip():
                any_text = True
            pages.append(
                IngestedPage(
                    page_index=i,
                    image_path=out,
                    text_layer=text or None,
                    width=width,
                    height=height,
                )
            )
    finally:
        pdf.close()

    kind = SourceKind.PDF_TEXT_LAYER if any_text else SourceKind.PDF_SCANNED
    return IngestedDocument(source_path=source_path, kind=kind, pages=pages)


# --- docx ----------------------------------------------------------------


def _ingest_docx(source_path: Path) -> IngestedDocument:
    import docx  # python-docx — lazy import

    doc = docx.Document(str(source_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(parts).strip()
    return IngestedDocument(
        source_path=source_path,
        kind=SourceKind.DOCX,
        pages=[IngestedPage(page_index=0, image_path=None, text_layer=text or None)],
    )
